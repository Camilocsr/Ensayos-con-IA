import openai
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv
import os

load_dotenv()
openai.api_key = os.getenv("APY_KEY_OPENIA")

def set_apa_style(doc):
    print("Aplicando estilo APA...")
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for paragraph in doc.paragraphs:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

def add_header(doc):
    print("Agregando encabezado...")
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "Ensayo sobre Programación"
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    header_paragraph.style.font.name = 'Times New Roman'
    header_paragraph.style.font.size = Pt(12)

def add_footer(doc):
    print("Agregando pie de página...")
    section = doc.sections[0]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = "Página "
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    page_run = footer_paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    page_run._r.append(fld_char1)
    page_run._r.append(instr_text)
    page_run._r.append(fld_char2)

    footer_paragraph.style.font.name = 'Times New Roman'
    footer_paragraph.style.font.size = Pt(12)

def corregir_ortografia_y_gramatica(texto):
    try:
        print("Corrigiendo ortografía y gramática...")
        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "Eres un experto en ortografía y gramática en español."},
                {"role": "user", "content": f"Por favor corrige la ortografía y gramática del siguiente texto:\n\n{texto}"}
            ]
        )
        texto_corregido = response['choices'][0]['message']['content'].strip()
        return texto_corregido
    except Exception as e:
        print(f"Error al corregir el texto: {e}")
        return texto

def leer_y_aplicar_apa(input_path, output_path):
    try:
        print("Leyendo documento original...")
        doc_original = Document(input_path)
        doc_nuevo = Document()

        for paragraph in doc_original.paragraphs:
            texto_corregido = corregir_ortografia_y_gramatica(paragraph.text)
            doc_nuevo.add_paragraph(texto_corregido)

        set_apa_style(doc_nuevo)
        add_header(doc_nuevo)
        add_footer(doc_nuevo)

        print(f"Guardando el documento modificado en {output_path}...")
        doc_nuevo.save(output_path)
        print(f"Se ha modificado el documento y guardado en {output_path}")
    except Exception as e:
        print(f"Error al procesar el documento: {e}")

if __name__ == "__main__":
    input_path = "documents/TRABAJODEINVESTIGACIÓN.docx"
    output_path = "documents/TRABAJODEINVESTIGACIÓN-corregido-con-ortografiaYNormasApa.docx"
    
    leer_y_aplicar_apa(input_path, output_path)