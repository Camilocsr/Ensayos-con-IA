import openai
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
import os

load_dotenv()

openai.api_key = os.getenv("APY_KEY_OPENIA")

def generar_ensayo_y_guardar(output_path):
    
    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "Eres un estudiante de primer semestre de ingeniería de software."},
            {"role": "user", "content": "Escribe un ensayo sobre programación de software en un estilo que parezca lo más humano posible, debe ser largo"}

        ]
    )
    
    
    ensayo = response['choices'][0]['message']['content'].strip()
    
    
    doc_nuevo = Document()
    
    
    titulo = doc_nuevo.add_heading("Ensayo sobre Programación", level=1)
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    
    doc_nuevo.add_paragraph(ensayo)
    
    doc_nuevo.save(output_path)
    
    print(f"Se ha generado el ensayo y guardado en {output_path}")

if __name__ == "__main__":
    output_path = "documents\ensayo.docx"
    generar_ensayo_y_guardar(output_path)